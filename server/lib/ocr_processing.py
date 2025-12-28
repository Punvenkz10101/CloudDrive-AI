#!/usr/bin/env python3
"""
OCR Processing using Tesseract OCR
Automatically extracts text from PDFs and images
"""
import os
import sys
import json
import logging
from pathlib import Path
from datetime import datetime
from pdf2image import convert_from_path
from PIL import Image, ImageEnhance
import pytesseract

# Configure Tesseract path for Windows
if sys.platform == 'win32':
    # Common Tesseract installation paths on Windows
    tesseract_paths = [
        r'C:\Program Files\Tesseract-OCR\tesseract.exe',
        r'C:\Program Files (x86)\Tesseract-OCR\tesseract.exe',
        r'C:\Tesseract-OCR\tesseract.exe',
    ]
    
    # Try to find Tesseract
    tesseract_found = False
    for tesseract_path in tesseract_paths:
        if os.path.exists(tesseract_path):
            pytesseract.pytesseract.tesseract_cmd = tesseract_path
            tesseract_found = True
            logging.info(f"Using Tesseract at: {tesseract_path}")
            break
    
    if not tesseract_found:
        # Try using tesseract from PATH (if installed via conda/choco)
        try:
            import shutil
            tesseract_cmd = shutil.which('tesseract')
            if tesseract_cmd:
                pytesseract.pytesseract.tesseract_cmd = tesseract_cmd
                logging.info(f"Using Tesseract from PATH: {tesseract_cmd}")
            else:
                logging.warning("Tesseract not found. Please install Tesseract OCR or add it to PATH.")
        except Exception:
            logging.warning("Tesseract not found. Please install Tesseract OCR or add it to PATH.")

# Configure logging (move before Tesseract config to see warnings)
logging.basicConfig(level=logging.INFO, format='%(levelname)s: %(message)s')
logger = logging.getLogger(__name__)

# Optional OpenCV for advanced preprocessing
try:
    import cv2
    import numpy as np
    OPENCV_AVAILABLE = True
except ImportError:
    OPENCV_AVAILABLE = False
    logger.warning("OpenCV not available - using basic image preprocessing")


class OCRProcessor:
    def __init__(self, output_dir=None):
        """Initialize OCR processor"""
        if output_dir is None:
            # Default to extracted_text directory relative to lib
            lib_dir = Path(__file__).parent
            output_dir = lib_dir.parent / 'extracted_text'
        self.output_dir = Path(output_dir)
        self.output_dir.mkdir(parents=True, exist_ok=True)
        logger.info(f"OCR output directory: {self.output_dir}")
        self.metadata_file = self.output_dir / 'metadata.json'
        self.metadata = self._load_metadata()
        logger.info(f"Loaded {len(self.metadata)} files from OCR metadata")
        
    def _load_metadata(self):
        """Load existing metadata"""
        if self.metadata_file.exists():
            try:
                with open(self.metadata_file, 'r', encoding='utf-8') as f:
                    return json.load(f)
            except Exception as e:
                logger.warning(f"Error loading metadata: {e}")
        return {}
    
    def _save_metadata(self):
        """Save metadata to file"""
        try:
            with open(self.metadata_file, 'w', encoding='utf-8') as f:
                json.dump(self.metadata, f, indent=2)
        except Exception as e:
            logger.error(f"Error saving metadata: {e}")
    
    def _get_mime_type(self, file_name):
        """Get MIME type from file extension"""
        ext = Path(file_name).suffix.lower()
        mime_map = {
            '.pdf': 'application/pdf',
            '.jpg': 'image/jpeg',
            '.jpeg': 'image/jpeg',
            '.png': 'image/png',
            '.bmp': 'image/bmp',
            '.tiff': 'image/tiff',
            '.gif': 'image/gif',
            '.docx': 'application/vnd.openxmlformats-officedocument.wordprocessingml.document',
            '.doc': 'application/msword'
        }
        return mime_map.get(ext, 'application/octet-stream')
    
    def _generate_description(self, file_name, text):
        """Generate a description from file name and content"""
        description_parts = []
        name_lower = file_name.lower()
        
        # Check for common document types
        if 'rsm' in name_lower or 'road' in name_lower or 'safety' in name_lower:
            description_parts.append("Road Safety and Management")
        if 'question' in name_lower or 'bank' in name_lower:
            description_parts.append("Question Bank")
        if 'pdf' in name_lower:
            description_parts.append("PDF Document")
        if any(ext in name_lower for ext in ['.jpg', '.jpeg', '.png']):
            description_parts.append("Image Document")
            
        # Check content
        content_preview = text[:200].lower() if text else ''
        if any(term in content_preview for term in ['question', 'answer', 'mcq', 'multiple choice']):
            description_parts.append("Contains Questions")
        if 'id' in content_preview or 'card' in content_preview or 'registration' in content_preview:
            description_parts.append("ID or Registration Document")
            
        if description_parts:
            return " - ".join(description_parts)
        return f"Document: {Path(file_name).stem}"
    
    def _preprocess_image(self, image):
        """Enhance image quality for better OCR accuracy"""
        try:
            if OPENCV_AVAILABLE:
                # Convert PIL to OpenCV format
                img_array = np.array(image.convert('RGB'))
                img_cv = cv2.cvtColor(img_array, cv2.COLOR_RGB2BGR)
                
                # Convert to grayscale
                gray = cv2.cvtColor(img_cv, cv2.COLOR_BGR2GRAY)
                
                # Denoise
                denoised = cv2.fastNlMeansDenoising(gray, None, 10, 7, 21)
                
                # Adaptive thresholding
                thresh = cv2.adaptiveThreshold(
                    denoised, 255, cv2.ADAPTIVE_THRESH_GAUSSIAN_C, 
                    cv2.THRESH_BINARY, 11, 2
                )
                
                # Convert back to PIL
                processed_image = Image.fromarray(thresh)
            else:
                # Basic grayscale conversion
                processed_image = image.convert('L')
            
            # Enhance contrast and sharpness
            enhancer = ImageEnhance.Contrast(processed_image)
            processed_image = enhancer.enhance(1.5)
            enhancer = ImageEnhance.Sharpness(processed_image)
            processed_image = enhancer.enhance(2.0)
            
            return processed_image
        except Exception as e:
            logger.warning(f"Image preprocessing failed, using original: {str(e)}")
            return image
    
    def extract_from_pdf(self, pdf_path):
        """Extract text from PDF using OCR with enhanced preprocessing"""
        try:
            logger.info(f"Processing PDF: {pdf_path}")
            # Use high DPI for better quality
            images = convert_from_path(pdf_path, dpi=300)
            full_text = ""
            
            # Optimized Tesseract config
            custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
            
            for i, image in enumerate(images):
                logger.info(f"Extracting text from page {i+1}/{len(images)}")
                processed_image = self._preprocess_image(image)
                text = pytesseract.image_to_string(
                    processed_image, 
                    config=custom_config, 
                    lang='eng'
                )
                full_text += f"\n--- Page {i+1} ---\n{text}"
            
            return full_text.strip()
        except Exception as e:
            logger.error(f"Error processing PDF {pdf_path}: {str(e)}")
            raise
    
    def extract_from_image(self, image_path):
        """Extract text from image using OCR with enhanced preprocessing"""
        try:
            logger.info(f"Processing image: {image_path}")
            
            # Verify file exists and is readable
            if not os.path.exists(image_path):
                raise FileNotFoundError(f"Image file not found: {image_path}")
            
            # Try to open and verify it's a valid image
            try:
                image = Image.open(image_path)
                # Verify the image by loading it
                image.verify()
            except Exception as img_error:
                # If verify() fails, try opening again (verify() closes the file)
                try:
                    image = Image.open(image_path)
                except Exception as img_error2:
                    raise ValueError(f"Cannot identify image file '{image_path}'. The file may be corrupted or not a valid image format. Error: {str(img_error2)}")
            
            # Re-open the image for processing (verify() closes it)
            image = Image.open(image_path)
            processed_image = self._preprocess_image(image)
            
            # Optimized Tesseract config
            custom_config = r'--oem 3 --psm 6 -c preserve_interword_spaces=1'
            text = pytesseract.image_to_string(
                processed_image, 
                config=custom_config, 
                lang='eng'
            )
            
            return text.strip()
        except FileNotFoundError:
            raise
        except ValueError as ve:
            logger.error(f"Invalid image file {image_path}: {str(ve)}")
            raise
        except Exception as e:
            logger.error(f"Error processing image {image_path}: {str(e)}")
            raise
    
    def extract_from_word(self, file_path):
        """Extract text from Word document (.docx)"""
        try:
            logger.info(f"Processing Word document: {file_path}")
            
            # Try python-docx
            try:
                from docx import Document
                doc = Document(file_path)
                full_text = []
                
                for para in doc.paragraphs:
                    full_text.append(para.text)
                
                for table in doc.tables:
                    for row in table.rows:
                        row_text = []
                        for cell in row.cells:
                            row_text.append(cell.text)
                        full_text.append('\t'.join(row_text))
                
                return '\n'.join(full_text).strip()
            except ImportError:
                # python-docx not available, try to convert to PDF and use OCR
                logger.warning("python-docx not available. Install with: pip install python-docx")
                logger.info("Attempting to extract using basic file reading...")
                
                # For .docx files, they're zip archives with XML
                try:
                    import zipfile
                    with zipfile.ZipFile(file_path, 'r') as zip_ref:
                        # Read the main document XML
                        xml_content = zip_ref.read('word/document.xml')
                        
                        # Simple text extraction from XML (basic approach)
                        import re
                        # Remove XML tags and extract text
                        text = re.sub(r'<[^>]+>', '\n', xml_content.decode('utf-8', errors='ignore'))
                        # Clean up whitespace
                        text = re.sub(r'\n+', '\n', text)
                        text = re.sub(r' +', ' ', text)
                        
                        return text.strip()
                except Exception as zip_error:
                    logger.error(f"Could not extract from Word document: {str(zip_error)}")
                    raise ValueError("Word document processing failed. Please install python-docx: pip install python-docx")
        except Exception as e:
            logger.error(f"Error processing Word document {file_path}: {str(e)}")
            raise
    
    def process_file(self, file_path, file_name=None):
        """
        Process any supported file (PDF or image) and extract text
        
        Args:
            file_path: Path to the file to process
            file_name: Optional name to use in metadata (should match storage blob name)
        
        Returns:
            Extracted text content
        """
        file_path = Path(file_path)
        file_ext = file_path.suffix.lower()
        
        # Use provided file_name or default to file_path name
        metadata_file_name = file_name if file_name else file_path.name
        
        if not file_path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")
        
        if file_ext == '.pdf':
            text = self.extract_from_pdf(file_path)
        elif file_ext in ['.jpg', '.jpeg', '.png', '.bmp', '.tiff', '.gif']:
            text = self.extract_from_image(file_path)
        elif file_ext in ['.docx', '.doc']:
            text = self.extract_from_word(file_path)
        else:
            raise ValueError(f"Unsupported file format: {file_ext}")
        
        # Save extracted text to file
        output_file = self.output_dir / f"{file_path.stem}_extracted.txt"
        with open(output_file, 'w', encoding='utf-8') as f:
            f.write(text)
        
        # Generate description
        description = self._generate_description(metadata_file_name, text)
        
        # Update internal metadata
        self.metadata[metadata_file_name] = {
            "extracted_file": str(output_file),
            "file_size": file_path.stat().st_size if file_path.exists() else 0,
            "processed_at": datetime.now().isoformat(),
            "text_length": len(text),
            "content": text[:15000],  # Store first 15000 chars
            "description": description,
            "has_full_content": True
        }
        self._save_metadata()
        
        # Sync to main metadata file used by search
        self._sync_to_main_metadata(metadata_file_name, text, description, output_file)
        
        logger.info(f"Successfully processed {metadata_file_name} - extracted {len(text)} characters")
        return text
    
    def _sync_to_main_metadata(self, file_name, content, description, output_file):
        """Sync OCR metadata to the main metadata file used by search"""
        try:
            lib_dir = Path(__file__).parent
            main_metadata_file = lib_dir.parent / 'storage' / 'file-metadata.json'
            main_metadata_file.parent.mkdir(parents=True, exist_ok=True)
            
            logger.info(f"Syncing to main metadata: {main_metadata_file}")
            
            # Read existing metadata
            if main_metadata_file.exists():
                try:
                    with open(main_metadata_file, 'r', encoding='utf-8') as f:
                        main_metadata = json.load(f)
                except json.JSONDecodeError:
                    logger.warning("Metadata file corrupted, creating new one")
                    main_metadata = {}
            else:
                main_metadata = {}
            
            # Ensure we use absolute path for fullContentPath
            if not Path(output_file).is_absolute():
                full_content_path = str(self.output_dir / Path(output_file).name)
            else:
                full_content_path = str(output_file)
            
            # Ensure path uses forward slashes for consistency
            full_content_path = full_content_path.replace('\\', '/')
            
            # Store more content in metadata (up to 15000 chars)
            content_preview = content[:15000] if len(content) > 15000 else content
            
            # Update main metadata
            main_metadata[file_name] = {
                "description": description,
                "content": content_preview,
                "mimeType": self._get_mime_type(file_name),
                "updatedAt": datetime.now().isoformat(),
                "hasFullContent": True,
                "fullContentPath": full_content_path,
                "textLength": len(content)
            }
            
            # Write back with error handling
            try:
                with open(main_metadata_file, 'w', encoding='utf-8') as f:
                    json.dump(main_metadata, f, indent=2, ensure_ascii=False)
                logger.info(f"✓ Synced metadata for {file_name} - {len(content)} chars, saved {len(content_preview)} to metadata")
            except Exception as write_error:
                logger.error(f"Failed to write metadata: {write_error}")
                raise
            
        except Exception as e:
            logger.error(f"Failed to sync to main metadata: {str(e)}")
            import traceback
            logger.error(traceback.format_exc())
            raise  # Re-raise to surface the error


if __name__ == '__main__':
    if len(sys.argv) < 3:
        print("Usage: python ocr_processing.py <file_path> <file_name>")
        sys.exit(1)
    
    file_path = sys.argv[1]
    file_name = sys.argv[2]
    
    processor = OCRProcessor()
    text = processor.process_file(file_path, file_name)
    print(f"SUCCESS: Extracted {len(text)} characters")

